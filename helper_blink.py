'''Helper module to be imported by main.py which scrapes all books in library
of service blinkist.com
'''
import os
import time
import msvcrt
import argparse
import tkinter as tk
from tkinter import filedialog
import random
#import re
from docx2pdf import convert

from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import NoSuchElementException

import docx
from docx import Document
#from docx.shared import Inches
#from docx.enum.text import WD_LINE_SPACING

#Code not to be executed on module import

def set_geckodriver_path():
    '''Point to Chrome webdriver local path'''
    print('\nPlease select the local geckodriver.exe\n')
    time.sleep(1.5)
    root = tk.Tk()
    root.withdraw()
    geckodriver_path = filedialog.askopenfilename()
    if geckodriver_path == '':
        geckodriver_path = 'geckodriver.exe'
    return geckodriver_path

def setup_geckodriver(): #normally with argument 'chrome_webdriver_path' but seems webdriver doesn't run in mopdule-function so transferred to main.py
    '''Setup of Chrome webdriver'''
    firefox_profile = webdriver.FirefoxProfile()
    firefox_profile.set_preference("browser.privatebrowsing.autostart", True)
    return firefox_profile

def get_blink_login():
    parser = argparse.ArgumentParser()
    parser.add_argument("usr", nargs= '?', default='', help="Provide your Blinkist username for automated login", type=str)
    parser.add_argument("pwd", nargs= '?', default='', help="Provide your Blinkist password for automated login", type=str)
    args = parser.parse_args()
    blink_pwd = args.pwd
    blink_usr = args.usr
        
    return blink_usr, blink_pwd

def blinkist_login(driver):
    '''Blinkist Login, with manual reCaptcha resolution'''
    blink_usr, blink_pwd = get_blink_login()
    driver.get('https://www.blinkist.com/en/nc/login')

    if blink_usr != '' and blink_pwd != '':
        
        time.sleep(random.uniform(0, 1.5))
        driver.find_element_by_id('login-form_login_email').send_keys(blink_usr)
        time.sleep(random.uniform(0, 0.5))
        driver.find_element_by_id('login-form_login_password').send_keys(blink_pwd)
        time.sleep(random.uniform(0, 0.5))
        driver.find_element_by_id('login-form_login_password').send_keys(Keys.TAB + Keys.TAB + \
            Keys.TAB + Keys.TAB + Keys.TAB + Keys.TAB + Keys.SPACE)
        print('\nPlease solve the "reCaptcha" and click "Login" on the blinkist website.\nOnce you have successfully logged in, press any key in this terminal.\n')

    else:

        print('You haven\'t specified your Blinkist login information. Please manually log in once the login screen is in focus.\nOnce you have successfully logged in, press any key in this terminal.\n')
    
    msvcrt.getch()

def usr_selection_books():
    '''Query user for selection of books to retrieve
    '''
    all_books = 'na'

    while all_books == 'na':
        print('Which book summaries do you want to retrieve?\n\
        1. The summaries of all the books in my library\n\
        2. The summaries of all the books there are on Blinkist (this will take some time and might get your IP blocked)')

        usr_selection = msvcrt.getch()
        if usr_selection == b'1':
            all_books = 1
            break
        if usr_selection == b'2':
            all_books = 0
            break

        print('This is not a valid input. Try again or press "Ctrl + C" to end the program.')
        all_books = 'na'

    return all_books


def get_read_urls_lib(driver):
    '''Get urls for all books in library'''
    driver.get('https://www.blinkist.com/en/nc/library')
    #time.sleep(random.uniform(1.1, 1.6))
    while True:
        try:
            driver.find_element_by_link_text('Load more').click()
            print('"Load more" clicked')
            time.sleep(random.uniform(1.5, 3.5))

        except NoSuchElementException:
            print('No element "Load more"')
            break

    read_urls = []
    #elems = driver.find_elements_by_xpath("//a[@href]")
    elems = driver.find_elements_by_css_selector('a.book-card__main-link')

    for elem in elems:
        r_url = elem.get_attribute("href")
        read_urls.append(r_url)
    print('\nList of book reader urls:')
    print(read_urls)
    print('Press any key to continue.')
    msvcrt.getch()
    print('')
    return read_urls

# def skip_already_processed_read_urls(read_urls, save_dir):
#     ''' Checks provided save directory for already processed books and removes them from the book list to be processed.'''
#     # Get titles of books already in the provided save directory and convert them to book read urls
#     files_in_dir = []
#     for files in os.listdir(save_dir):
#         if files.endswith('.docx'):
#             # split string at the last occurence of '-'
#             # delete special characters (including '-')
#             # create the actual link to compare with the url-List to be scraped
#             files = files.rpartition('-')[0]
#             files = re.sub('[^A-Za-z0-9 ]+', '', files)
#             files_in_dir.append('https://www.blinkist.com/en/nc/reader/' + str(files.lower().replace(' ', '-')) + '-en')
#     # Compare list with already processed urls with list of book urls to be processed, and delete duplicates
#     number_read_urls_pre = len(read_urls)
#     for url in read_urls[:]:
#         if url in files_in_dir:
#             read_urls.remove(url)
#     skipped_urls = number_read_urls_pre - len(read_urls)

#     print('\n# of read urls before skipping: ' + str(number_read_urls_pre))
#     print('\n# of read urls after skipping: ' + str(len(read_urls)) + '\n')
#     print(str(skipped_urls) + ' urls have been removed because the respective documents already exist in the provided directory.\nPress "Enter" to continue.\n')

#     msvcrt.getch()
#     return read_urls, skipped_urls

def skip_already_processed_read_urls(read_urls, save_dir):
    ''' Checks provided save directory for already processed books and removes them from the book list to be processed.'''
    # Extract url from docx and added to comparison url list
    files_in_dir = []
    for file in os.listdir(save_dir):
        if file.endswith('.docx'):
            doc = docx.Document(save_dir + '\\' + file)
            all_paras = doc.paragraphs
            for para in all_paras:
                if 'https://www.blinkist' in para.text:
                    files_in_dir.append(para.text.split(': ')[1])

    # Compare list with already processed urls with list of book urls to be processed, and delete duplicates
    number_read_urls_pre = len(read_urls)
    for url in read_urls[:]:
        if url in files_in_dir:
            read_urls.remove(url)
    skipped_urls = number_read_urls_pre - len(read_urls)

    print('\n# of read urls before skipping: ' + str(number_read_urls_pre))
    print('\n# of read urls after skipping: ' + str(len(read_urls)) + '\n')
    print(str(skipped_urls) + ' urls have been removed because the respective documents already exist in the provided directory.\nPress "Enter" to continue.\n')

    msvcrt.getch()
    return read_urls, skipped_urls

def get_book_urls(read_urls):
    '''Get read urls for all book from list of book urls'''
    book_urls = []
    for elem in read_urls:
        new_link = elem.replace('/en/nc/reader', '/en/books')
        book_urls.append(new_link)
    print('\nList of book urls:\n' + str(book_urls))
    return book_urls

def get_save_dir():
    '''Query user for save directory of final doc'''
    root = tk.Tk()
    root.withdraw()
    print('Provide a directory to store the book summaries (the default directory is "blinkist_summaries").')
    save_dir = filedialog.askdirectory()
    if save_dir == '':
        if not os.path.exists('blinkist_summaries'):
            os.mkdir('blinkist_summaries')
            print('Directory "blinkist_summaries" created as standard save directory!')
        save_dir = 'blinkist_summaries'
    print('\nYou chose ' + str(save_dir) +' as directory to store the books summaries.\n')
    return save_dir

def get_book_info(book_url, driver):
    '''Collect book info'''
    driver.get(book_url)
    time.sleep(random.uniform(1, 1.5))

    #title
    header_title = (driver.find_element_by_class_name('book__header__title')).text
    print(header_title + '\n')

    #subtitle
    header_subtitle = (driver.find_element_by_class_name('book__header__subtitle')).text
    #header_subtitle = re.sub('[^A-Za-z0-9_ -]+', '', header_subtitle)
    print(header_subtitle + '\n')

    #author
    header_author = (driver.find_element_by_class_name('book__header__author')).text
    if ',' in header_author:
        header_author = header_author.split(',')[0]
    print(header_author + '\n')


    #readtime
    for item in driver.find_elements_by_class_name('book__header__info'):
        read_time = (item.text).split('\n')[0] + '\n'
        print(read_time)

    #summary
    for element in driver.find_elements_by_tag_name('p'):
        if element.text != '':
            book_info = element.text + '\n'
            print(book_info)
    #amazon link
    elems = driver.find_elements_by_xpath("//a[@href]")
    for elem in elems:
        if 'amazon' in elem.get_attribute("href"):
            _amazon_link = elem.get_attribute("href")

    try:
        amazon_link = _amazon_link.rsplit('/', 1)[0]
    except UnboundLocalError:
        print('UnboundLocalError occured...')
        amazon_link = 'Amazon url not retrievable'
    #amazon_link = _amazon_link.rsplit('/', 1)[0]
    print(amazon_link)

    return (header_title, header_subtitle, header_author, read_time, book_info, amazon_link)


def get_book_headers_and_chapters(read_url, driver):
    '''Collect headers and chapters of iterated book'''
    driver.get(read_url)
    time.sleep(random.uniform(1, 1.5))

    # Make all needed elements visible
    container = driver.find_elements_by_css_selector('div.chapter')
    for elem in container:
        driver.execute_script("arguments[0].style.display = 'block'; arguments[0].style.opacity = '1'", elem)
    # container = driver.find_elements_by_xpath('/html/body/div[3]/main/div[1]/div[3]/article/div')
    # for elem in container:
    #     driver.execute_script("arguments[0].style.display = 'block';", elem) ###SOMETHING IS WRONG WITH THIS LINE!!!

    # Get all chapter headers and number of chapter headers of book
    chapter_headers_raw = driver.find_elements_by_css_selector('div.chapter > h1')
    # chapter_headers_raw = driver.find_elements_by_css_selector('body > div.page > main > div.reader__container > div.reader__container__right > article > div > h1')
    num_of_chapters = len(chapter_headers_raw)
    chapter_headers = []
    for header in chapter_headers_raw:
        chapter_headers.append(header.text)

    # Get all chapter body texts. Number of body texts includes an additional useless body text in the end
    # which is why we refer to the number of headers

    chapter_bodies_raw = driver.find_elements_by_css_selector('div.chapter > div')
    chapter_bodies = []
    for body in chapter_bodies_raw:
        chapter_bodies.append(body.text)
    print('\n Number of chapters, chapter headers, chapter bodies')
    print(chapter_headers)
    print(num_of_chapters)
    print(chapter_bodies)

    return chapter_headers, chapter_bodies, num_of_chapters


def file_name(header_title, header_author):
    ''' Create file name for doc'''
    book_fn = header_title + '-' + header_author.replace(' ', '_')[3:]

    # Account for invalid characters in Windows and replace with valid ones
    invalid_chars = [':', 'â€“', '/', '?', '*']
    valid_chars = [';', '—', '~', '�', '★']
    for i, _ in enumerate(invalid_chars):
        book_fn = book_fn.replace(invalid_chars[i], valid_chars[i])

    return book_fn


def create_doc(header_title, header_subtitle, header_author, read_time, book_info, amazon_link, chapter_headers, chapter_bodies, num_of_chapters, save_dir, book_fn, read_url):
    '''Create doc for iterated book'''
    document = Document()

    document.add_heading(header_title, 0)
    document.add_heading(header_subtitle, 1)
    document.add_heading(header_author, 3)
    document.add_heading(read_time, 8)
    paragraph = document.add_paragraph(book_info)
    paragraph.paragraph_format.line_spacing = 1.5
    document.add_paragraph('Book on Amazon: ' + str(amazon_link))
    document.add_paragraph('Book on Blinkist: ' + str(read_url))

    for index in range(num_of_chapters):
        document.add_page_break()
        document.add_heading(chapter_headers[index], 2)
        paragraph = document.add_paragraph('\n' + chapter_bodies[index])
        paragraph.paragraph_format.line_spacing = 1.5

    document.save(str(save_dir) + '/' + str(book_fn) + '.docx')
    print('Created: ' + str(save_dir) + '/' + str(book_fn) + '.docx')

def convert_docx_to_pdf(save_dir):
    ''' Convert all docx files in folder to pdf'''
    convert(save_dir)
